# HONGXUN-LOCKED — 此文件受保护，修改需密码验证
# 请通过程序界面解锁后编辑
"""
鸿讯 HONGXUN · 论文检索引擎 — 编排协调层
版本 1.0.0
"""

from datetime import datetime, timedelta
import os
import io
import json
import subprocess
import sys
import time
import threading
from contextlib import contextmanager

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config_manager import OUTPUT_DIR, UNSENT_DIR, load_push_records, add_push_record, load_app_config
from .search import search_papers
from .abstract import enrich_abstract
from .email_sender import send_email

# 全局取消标志（用于中止正在执行的检索任务）
_search_cancel_flag = threading.Event()


def cancel_current_search():
    """设置取消标志，通知正在执行的检索任务停止"""
    _search_cancel_flag.set()


def reset_search_cancel():
    """清除取消标志"""
    _search_cancel_flag.clear()


def is_search_cancelled() -> bool:
    """检查是否有取消请求"""
    return _search_cancel_flag.is_set()


# ── macOS 防休眠（使用 caffeinate） ─────────────────────────

def _caffeinate_process():
    """启动 macOS caffeinate 进程防止系统休眠"""
    if sys.platform != "darwin":
        return None
    try:
        proc = subprocess.Popen(
            ["caffeinate", "-dimsu"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        return proc
    except Exception:
        return None


def _stop_caffeinate(proc):
    """停止 caffeinate 进程"""
    if proc is not None:
        try:
            proc.terminate()
            proc.wait(timeout=3)
        except Exception:
            try:
                proc.kill()
            except Exception:
                pass


@contextmanager
def keep_awake():
    """上下文管理器：在代码块执行期间阻止 macOS 休眠"""
    proc = _caffeinate_process()
    try:
        yield
    finally:
        _stop_caffeinate(proc)


# ── 辅助样式 ──────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str):
    """设置表格单元格底色"""
    from docx.oxml.ns import qn
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def _set_run_font(run, text: str, size=Pt(10.5), bold=False):
    """为 run 设置西文(Times New Roman)和中文(宋体)字体"""
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.bold = bold
    from docx.oxml.ns import qn
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'SimSun')


def _set_paragraph_spacing(paragraph, line_spacing=Pt(18)):
    """设置固定行间距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _init_docx_styles(doc):
    """一次性初始化文档全局样式，避免每篇论文重复操作"""
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = Pt(18)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Pt(0)

    # 设置中文字体（SimSun）到 Normal 样式级，一次生效
    from docx.oxml.ns import qn
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = style.element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'SimSun')


def _style_cell_key(cell, color='F8FAFC'):
    """快速设置标签列样式（仅加粗，不重复字体/字号）"""
    _set_cell_shading(cell, color)
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True


def _build_docx(papers, task_data, title_prefix="", progress_callback=None) -> Document:
    """
    根据论文列表构建格式规范的 DOCX 报告。
    字体：西文 Times New Roman，中文 宋体，5号(10.5pt)，不加粗
    行间距：固定 18 磅
    摘要：首行缩进 2 字符 (21pt)
    """
    doc = Document()

    # ── 全局样式（一次性） ──
    _init_docx_styles(doc)

    # ── 标题 ──
    title = doc.add_heading(f'{title_prefix}{task_data["name"]} - 论文检索报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 元信息表 ──
    meta = doc.add_table(rows=5, cols=2, style='Table Grid')
    meta_data = [
        ("检索时间", datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ("时间范围", f"{task_data['date_start']} 至 {task_data['date_end']}"),
        ("关键词", '; '.join(task_data['keywords'])),
        ("期刊", '; '.join(task_data['journals'])),
        ("命中总数", f"{len(papers)} 篇"),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
    _style_cell_key(meta.cell(0, 0), 'EFF6FF')
    _style_cell_key(meta.cell(1, 0), 'EFF6FF')
    _style_cell_key(meta.cell(2, 0), 'EFF6FF')
    _style_cell_key(meta.cell(3, 0), 'EFF6FF')
    _style_cell_key(meta.cell(4, 0), 'EFF6FF')

    doc.add_paragraph()

    # ── 逐篇论文 ──
    for idx, p in enumerate(papers, 1):
        # 先更新进度，再处理该篇
        if progress_callback:
            msg = f"生成报告 ({idx}/{len(papers)})"
            progress_callback(0.90 + (idx / len(papers)) * 0.09, msg)

        kws = ', '.join(p.get('matched_keywords', []))

        doc.add_heading(f'{idx}. {p["title"]}', level=2)

        # 使用文本行代替表格（大幅减少 XML 体积，避免大规模文档卡死）
        field_lines = [
            ("期刊", p.get('container_title', '')),
            ("作者", p['authors']),
            ("发表时间", p['pub_date']),
            ("DOI", p['doi']),
            ("匹配关键词", kws),
        ]
        title_cn = p.get('title_cn', '')
        if title_cn and isinstance(title_cn, str):
            field_lines.append(("中文标题", title_cn))
        affs = p.get('affiliations', {})
        if affs:
            aff_text = '；'.join(f"{a}: {o}" for a, o in affs.items())
            field_lines.append(("机构", aff_text))

        for label, value in field_lines:
            if not isinstance(value, str):
                value = str(value) if value is not None else ''
            line = doc.add_paragraph(f'【{label}】{value}')
            _set_paragraph_spacing(line)

        abstract = p['abstract']
        abstract_cn = p.get('abstract_cn', '')
        if not isinstance(abstract_cn, str):
            abstract_cn = str(abstract_cn) if abstract_cn else ''
        if abstract and abstract != '无摘要':
            if not isinstance(abstract, str):
                abstract = str(abstract) if abstract else '无摘要'
            ab_p = doc.add_paragraph(f'摘要：{abstract}')
            ab_p.paragraph_format.first_line_indent = Pt(21)
            if abstract_cn:
                ab_cn_p = doc.add_paragraph(f'中文摘要：{abstract_cn}')
                ab_cn_p.paragraph_format.first_line_indent = Pt(21)
        else:
            doc.add_paragraph('摘要：无摘要')

    return doc


def _build_txt(papers, task_data, title_prefix="") -> str:
    """生成纯文本报告"""
    lines = []
    lines.append(f"{title_prefix}{task_data['name']} - 论文检索报告")
    lines.append("=" * 50)
    lines.append(f"检索时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"时间范围：{task_data['date_start']} 至 {task_data['date_end']}")
    lines.append(f"关键词：{'; '.join(task_data['keywords'])}")
    lines.append(f"期刊：{'; '.join(task_data['journals'])}")
    lines.append(f"命中总数：{len(papers)} 篇")
    lines.append("=" * 50)
    lines.append("")

    for idx, p in enumerate(papers, 1):
        kws = ', '.join(p.get('matched_keywords', []))
        title_cn = p.get('title_cn', '')
        abstract_cn = p.get('abstract_cn', '')
        lines.append(f"{idx}. {p['title']}")
        if title_cn:
            lines.append(f"   中文标题：{title_cn}")
        lines.append(f"   期刊：{p.get('container_title', '')}")
        lines.append(f"   作者：{p['authors']}")
        lines.append(f"   发表时间：{p['pub_date']}")
        lines.append(f"   DOI：{p['doi']}")
        lines.append(f"   匹配关键词：{kws}")
        abstract = p['abstract']
        if abstract and abstract != '无摘要':
            lines.append(f"   摘要：{abstract}")
            if abstract_cn:
                lines.append(f"   中文摘要：{abstract_cn}")
        else:
            lines.append(f"   摘要：无摘要")
        lines.append("")

    return '\n'.join(lines)


def _build_md(papers, task_data, title_prefix="", progress_callback=None) -> str:
    """生成 Markdown 报告"""
    lines = []
    lines.append(f"# {title_prefix}{task_data['name']} - 论文检索报告")
    lines.append("")
    lines.append(f"**检索时间**：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"**时间范围**：{task_data['date_start']} 至 {task_data['date_end']}")
    lines.append(f"**关键词**：{'; '.join(task_data['keywords'])}")
    lines.append(f"**期刊**：{'; '.join(task_data['journals'])}")
    lines.append(f"**命中总数**：{len(papers)} 篇")
    lines.append("")
    lines.append("---")
    lines.append("")

    for idx, p in enumerate(papers, 1):
        if progress_callback:
            progress_callback(0.90 + (idx / len(papers)) * 0.09,
                              f"生成报告 ({idx}/{len(papers)})")
        kws = ', '.join(p.get('matched_keywords', []))
        title_cn = p.get('title_cn', '')
        abstract_cn = p.get('abstract_cn', '')
        if not isinstance(title_cn, str):
            title_cn = str(title_cn) if title_cn else ''
        if not isinstance(abstract_cn, str):
            abstract_cn = str(abstract_cn) if abstract_cn else ''
        lines.append(f"## {idx}. {p['title']}")
        if title_cn:
            lines.append("")
            lines.append(f"> **中文标题**：{title_cn}")
        lines.append("")
        lines.append(f"- **期刊**：{p.get('container_title', '')}")
        lines.append(f"- **作者**：{p['authors']}")
        lines.append(f"- **发表时间**：{p['pub_date']}")
        lines.append(f"- **DOI**：{p['doi']}")
        lines.append(f"- **匹配关键词**：{kws}")
        abstract = p['abstract']
        if abstract and abstract != '无摘要':
            lines.append(f"- **摘要**：{abstract}")
            if abstract_cn:
                lines.append(f"- **中文摘要**：{abstract_cn}")
        else:
            lines.append(f"- **摘要**：无摘要")
        lines.append("")

    return '\n'.join(lines)


def build_report(papers, task_data, file_path: str, title_prefix="", progress_callback=None):
    """
    根据 file_path 扩展名自动选择输出格式并将报告写入文件。
    支持：.doc / .docx / .txt / .md
    """
    ext = os.path.splitext(file_path)[1].lower()

    # 大规模论文时，对非 docx 格式仍做 gc 友好处理
    if ext in ('.doc', '.docx'):
        doc = _build_docx(papers, task_data, title_prefix, progress_callback)
        doc.save(file_path)
    elif ext == '.txt':
        content = _build_txt(papers, task_data, title_prefix)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
    elif ext == '.md':
        content = _build_md(papers, task_data, title_prefix, progress_callback)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
    else:
        file_path += '.docx'
        doc = _build_docx(papers, task_data, title_prefix, progress_callback)
        doc.save(file_path)


# ── 历史检索 → 文件 ───────────────────────────────────────

def run_history_search(task_id: str, task_data: dict, file_path: str, progress_callback=None) -> str:
    """执行历史论文检索（CrossRef + 知网），保存为用户指定的文件，返回文件路径。

    关键词过滤顺序：
    1. CrossRef 检索（按期刊 + 时间范围，不按关键词过滤——所有论文先获取）
    2. 摘要补全
    3. 用补全后的标题+摘要匹配 keywords1（仅保留至少匹配一个关键词的论文）
    4. 翻译（仅在全部过滤完成后执行）

    时间范围自动按月度切割，逐月检索后合并，避免 cursor 翻页不稳定。
    """
    # 使用 keep_awake 防止电脑休眠
    with keep_awake():
        if progress_callback:
            progress_callback(0.0, "搜索论文中...")

        # ========== 按月切割时间范围，逐月检索 ==========
        from calendar import monthrange
        start_dt = datetime.strptime(task_data["date_start"], "%Y-%m-%d")
        end_dt = datetime.strptime(task_data["date_end"], "%Y-%m-%d")

        # 生成月份列表（每月 1 日）
        months = []
        cursor = start_dt.replace(day=1)
        while cursor <= end_dt:
            months.append(cursor)
            if cursor.month == 12:
                cursor = cursor.replace(year=cursor.year + 1, month=1)
            else:
                cursor = cursor.replace(month=cursor.month + 1)

        total_months = len(months)
        seen_dois = set()
        all_papers = []

        for mi, m_first in enumerate(months):
            # 检查取消请求
            if is_search_cancelled():
                reset_search_cancel()
                raise KeyboardInterrupt("检索已取消")

            # 当月最后一天
            _, last_day = monthrange(m_first.year, m_first.month)
            m_last = m_first.replace(day=last_day)
            if m_last > end_dt:
                m_last = end_dt

            m_start_str = m_first.strftime("%Y-%m-%d")
            m_end_str = m_last.strftime("%Y-%m-%d")

            if progress_callback:
                progress_callback(0.0 + (mi / total_months) * 0.18,
                                  f"检索 {m_start_str} ～ {m_end_str}（{mi+1}/{total_months}月）")

            month_papers = search_papers(
                task_data["journals"],
                m_start_str, m_end_str,
                keywords=task_data.get("keywords"),
                progress_callback=None,  # 月度检索内部不传进度，避免嵌套回调
            )

            # 月度去重合并
            for p in month_papers:
                doi = p.get("doi", "")
                key = doi or (p.get("container_title", "") + p.get("title", ""))
                if key and key not in seen_dois:
                    seen_dois.add(key)
                    all_papers.append(p)

        papers = all_papers

        if progress_callback:
            progress_callback(0.20, f"检索完成（{len(papers)}篇），补全摘要中...")

        papers = enrich_abstract(papers, progress_callback)

        # 检查取消请求
        if is_search_cancelled():
            reset_search_cancel()
            raise KeyboardInterrupt("检索已取消")

        # 用补全后的摘要重新匹配关键词（补全前的摘要可能较短/不全，导致匹配不准）
        # 仅在标题+摘要中至少匹配一个关键词的论文才被保留
        raw_keywords = task_data["keywords"]
        keywords_lower = [kw.lower() for kw in raw_keywords]
        filtered_papers = []
        for p in papers:
            title = p.get("title", "")
            abstract = p.get("abstract", "")
            text = f"{title} {abstract}".lower()
            matched_lower = [kw for kw in keywords_lower if kw in text]
            if not matched_lower:
                continue  # 摘要补全后仍不匹配关键词，丢弃
            p["matched_keywords"] = [
                raw_keywords[keywords_lower.index(kw)] for kw in matched_lower
            ]
            filtered_papers.append(p)
        papers = filtered_papers

        if progress_callback:
            progress_callback(0.45, f"关键词匹配完成，保留{len(papers)}篇")

        if progress_callback:
            progress_callback(0.45, "准备生成报告...")

        if progress_callback:
            progress_callback(0.90, "生成报告中...")

        build_report(papers, task_data, file_path, progress_callback=progress_callback)

        if progress_callback:
            progress_callback(1.0, "完成")
        return file_path


# ── 增量检查 → 返回新论文 ────────────────────────────────

def run_increment_check(task_id: str, task_data: dict, start_str: str = None, end_str: str = None) -> list[dict]:
    """
    执行增量检查，返回新增论文列表（每篇附带 task_id, task_name）。
    由调度器统一发送邮件。

    每日推送仅按期刊检索，不按关键词过滤——获取期刊所有论文后
    补全摘要即发送。关键词仅用于历史检索。

    start_str / end_str: 可选自定义时间范围，默认向前 24 小时。
    """
    end_dt = datetime.now()
    if end_str:
        try:
            end_dt = datetime.strptime(end_str.split(" ")[0], "%Y-%m-%d")
        except Exception:
            end_dt = datetime.now()
    if start_str:
        try:
            start_dt = datetime.strptime(start_str.split(" ")[0], "%Y-%m-%d")
        except Exception:
            start_dt = end_dt - timedelta(hours=24)
    else:
        start_dt = end_dt - timedelta(hours=24)

    start_str_fmt = start_dt.strftime("%Y-%m-%d")
    end_str_fmt = end_dt.strftime("%Y-%m-%d")

    # 每日推送——仅按期刊检索，不传 keywords，获取期刊所有论文
    papers = search_papers(
        task_data["journals"],
        start_str_fmt, end_str_fmt,
    )

    # 补全摘要，不按关键词过滤，所有论文均保留
    papers = enrich_abstract(papers)

    pushed = load_push_records().get(task_id, [])
    new_papers = [p for p in papers if p["doi"] not in pushed]

    if new_papers:
        # 这里不调用 add_push_record——由调用方在邮件发送成功后执行
        for p in new_papers:
            p["_task_id"] = task_id
            p["_task_name"] = task_data["name"]

    return new_papers


# ── 合并多任务结果 → 一封邮件 ────────────────────────────

# SMTP 供应商端口备选表
_SMTP_FALLBACK_PORTS = {
    "smtp.qq.com": ["465", "587"],
    "smtp.163.com": ["465", "587"],
    "smtp.126.com": ["465", "587"],
    "smtp.gmail.com": ["465", "587"],
    "smtp.sina.com": ["465"],
    "smtp-mail.outlook.com": ["587"],
    "smtp.189.cn": ["465"],
}


def _smtp_send(cfg: dict, msg, receivers: list[str]) -> None:
    """
    发送邮件，自动尝试供应商的备选端口。
    配置端口失败时自动尝试同服务器的其他端口。
    """
    import smtplib

    server = cfg["smtp_server"]
    port = cfg["port"]
    ports = _SMTP_FALLBACK_PORTS.get(server, [port])

    # 配置的端口优先
    if port in ports:
        ports.remove(port)
    ports.insert(0, port)

    last_error = None
    for p in dict.fromkeys(ports):  # 去重保序
        try:
            p_int = int(p)
            if p_int == 465:
                s = smtplib.SMTP_SSL(server, p_int)
            else:
                s = smtplib.SMTP(server, p_int)
                s.starttls()
            s.login(cfg["sender"], cfg["auth_code"])
            s.sendmail(cfg["sender"], receivers, msg.as_string())
            s.quit()
            return
        except Exception as e:
            last_error = e
            continue
    if last_error:
        raise last_error


def send_combined_email(results: list[tuple[str, str, list[dict]]]) -> int:
    """
    将多个任务的增量结果合并为一封邮件发送。
    results: [(task_id, task_name, [新增论文]), ...]
    返回发送的邮件数（0 或 1）。
    """
    from email.mime.multipart import MIMEMultipart
    from email.mime.base import MIMEBase
    from email import encoders
    from email.mime.text import MIMEText

    from .config_manager import load_email_config
    from .coupon_manager import is_feature_allowed

    if not is_feature_allowed():
        return 0

    cfg = load_email_config()
    if not cfg.get("sender") or not cfg.get("auth_code"):
        return 0

    receivers = cfg.get("receivers", [])
    if isinstance(receivers, str):
        receivers = [r.strip() for r in receivers.replace('；', ';').split(';') if r.strip()]
    if not receivers and cfg.get("receiver", "").strip():
        receivers = [cfg["receiver"].strip()]
    if not receivers:
        return 0

    # 收集所有有新增论文的任务
    active_results = [(tid, tname, papers) for tid, tname, papers in results if papers]
    if not active_results:
        return 0

    total_papers = sum(len(p) for _, _, p in active_results)
    now_str = datetime.now().strftime('%Y%m%d')

    # ── 构建正文 ──
    body_parts = [f"论文更新提醒 — 共 {len(active_results)} 个任务，新增 {total_papers} 篇论文\n"]
    for tid, tname, papers in active_results:
        body_parts.append(f"▸ {tname}（{len(papers)} 篇）")
        for p in papers:
            kw = ', '.join(p.get('matched_keywords', []))
            body_parts.append(f"  · {p['title']}  [{p.get('container_title', '')}]  {kw}")
        body_parts.append("")
    body_parts.append(f"详情请查看各附件。")
    body_text = "\n".join(body_parts)

    msg = MIMEMultipart()
    msg['From'] = cfg["sender"]
    msg['To'] = '; '.join(receivers)
    msg['Subject'] = f"论文更新提醒 | {len(active_results)}个任务 | 新增{total_papers}篇"

    msg.attach(MIMEText(body_text, 'plain', 'utf-8'))

    # ── 先生成所有 DOCX 到本地，再构造附件 ──
    saved_files = []
    for tid, tname, papers in active_results:
        task_data_stub = {"name": tname, "date_start": "", "date_end": "",
                          "keywords": [], "journals": []}
        buf = io.BytesIO()
        doc = _build_docx(papers, task_data_stub, title_prefix=f"[增量] ")
        doc.save(buf)
        buf.seek(0)
        docx_bytes = buf.read()

        filename = f"{tname}_增量_{now_str}.docx"
        # 保存到本地
        os.makedirs(UNSENT_DIR, exist_ok=True)
        local_path = os.path.join(UNSENT_DIR, filename)
        with open(local_path, 'wb') as f:
            f.write(docx_bytes)
        saved_files.append(local_path)

        part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        part.set_payload(docx_bytes)
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment', filename=('utf-8', '', filename))
        msg.attach(part)

    try:
        _smtp_send(cfg, msg, receivers)
        # 发送成功 → 清理本地文件
        for fp in saved_files:
            try:
                os.remove(fp)
            except Exception:
                pass
        # 清理 pending_email.json（如果有）
        pending_path = os.path.join(UNSENT_DIR, "pending_email.json")
        if os.path.exists(pending_path):
            try:
                os.remove(pending_path)
            except Exception:
                pass
        return 1
    except Exception as e:
        print(f"合并邮件发送失败: {str(e)}", flush=True)
        # 发送失败 → 写入 pending_email.json 记录未发送文件
        try:
            pending = {
                "created_at": datetime.now().isoformat(),
                "error": str(e),
                "receivers": receivers,
                "subject": f"论文更新提醒 | {len(active_results)}个任务 | 新增{total_papers}篇",
                "files": saved_files,
                "active_tasks": [(tid, tname, len(papers)) for tid, tname, papers in active_results],
            }
            os.makedirs(UNSENT_DIR, exist_ok=True)
            with open(os.path.join(UNSENT_DIR, "pending_email.json"), 'w', encoding='utf-8') as f:
                json.dump(pending, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
        return 0


def _send_docx_email(subject: str, docx_bytes: bytes, filename: str) -> bool:
    """发送带 DOCX 附件的邮件"""
    from email.mime.multipart import MIMEMultipart
    from email.mime.base import MIMEBase
    from email import encoders

    from .config_manager import load_email_config
    from .coupon_manager import is_feature_allowed

    if not is_feature_allowed():
        return False

    cfg = load_email_config()
    if not cfg.get("sender") or not cfg.get("auth_code"):
        return False

    receivers = cfg.get("receivers", [])
    if isinstance(receivers, str):
        receivers = [r.strip() for r in receivers.replace('；', ';').split(';') if r.strip()]
    if not receivers and cfg.get("receiver", "").strip():
        receivers = [cfg["receiver"].strip()]
    if not receivers:
        return False

    msg = MIMEMultipart()
    msg['From'] = cfg["sender"]
    msg['To'] = '; '.join(receivers)
    msg['Subject'] = subject

    body = MIMEText(f'检测到新增论文，详情请查看附件「{filename}」', 'plain', 'utf-8')
    msg.attach(body)

    part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
    part.set_payload(docx_bytes)
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment', filename=('utf-8', '', filename))
    msg.attach(part)

    try:
        _smtp_send(cfg, msg, receivers)
        return True
    except Exception as e:
        print(f"邮件发送失败: {str(e)}", flush=True)
        return False
