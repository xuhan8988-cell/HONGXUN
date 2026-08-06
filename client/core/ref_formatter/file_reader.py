"""文件文本提取：把上传的格式要求文件（PDF/Word/文本）转为纯文本，供 LLM 解析。

PDF 优先用 markitdown（微软开源，GitHub 171k★，PDF/Word 等转 Markdown）；
未安装时降级 pdfminer。Word 用 python-docx，文本直接读。

安全：解析完成后调用方负责删除临时文件（数据不落盘）。
"""

import os


def extract_text_from_file(path: str) -> str:
    """按扩展名提取文件文本，返回字符串。失败返回空字符串。"""
    if not path or not os.path.exists(path):
        return ""
    ext = (os.path.splitext(path)[1] or "").lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext in (".docx", ".doc"):
        return _extract_docx(path)
    if ext in (".txt", ".md", ".markdown"):
        return _extract_text(path)
    # 未知类型：尝试按文本读
    return _extract_text(path)


def _extract_pdf(path: str) -> str:
    # 优先 markitdown（轻量，纯文本 PDF 转 Markdown 质量好）
    try:
        from markitdown import MarkItDown
        md = MarkItDown()
        result = md.convert(path)
        return (result.text_content or "").strip()
    except Exception:
        pass
    # 降级 pdfminer（已安装）
    try:
        from pdfminer.high_level import extract_text
        return (extract_text(path) or "").strip()
    except Exception:
        pass
    # 再降级 pymupdf4llm（若存在）
    try:
        import pymupdf4llm
        return (pymupdf4llm.to_markdown(path) or "").strip()
    except Exception:
        return ""


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if (p.text or "").strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_text(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    except Exception:
        return ""


def truncate_content(text: str, max_chars: int = 8000) -> str:
    """截断文本到 max_chars 字符（约 8000 token），保留首尾。"""
    if not text:
        return ""
    if len(text) <= max_chars:
        return text
    head = text[: int(max_chars * 0.8)]
    tail = text[-int(max_chars * 0.2):]
    return head + "\n\n……（中间内容已截断）……\n\n" + tail


def human_size(path: str) -> str:
    """返回文件大小的人类可读字符串，如 2.3MB。"""
    try:
        size = os.path.getsize(path)
    except Exception:
        return "0B"
    for unit in ("B", "KB", "MB", "GB"):
        if size < 1024:
            return f"{size:.0f}{unit}" if unit == "B" else f"{size:.1f}{unit}"
        size /= 1024
    return f"{size:.1f}GB"
