"""交叉引用：书签 + 内部超链接 + 角标上标 + 连续引用合并。"""

from docx import Document

from .utils.docx_utils import (
    add_bookmark_to_paragraph,
    replace_citation_with_hyperlink,
    replace_citation_text_in_paragraph,
    extract_reference_paragraphs,
    find_references_heading_index,
)


class CrossReferenceManager:
    """在 Word 文档中生成交叉引用超链接。"""

    def apply(self, input_path: str, output_path: str, formatted_refs: list[str],
              doc_data, superscript: bool = True, order_map: dict = None) -> dict:
        """写回格式化参考文献 + 加书签 + 正文引用超链接。

        order_map: 若做过重排，把正文中旧编号文本也重编号。
        返回统计 {hyperlinks_added, bookmarks_added, citations_renumbered}。
        """
        doc = Document(input_path)
        stats = {"hyperlinks_added": 0, "bookmarks_added": 0, "citations_renumbered": 0}

        idx = find_references_heading_index(doc)
        ref_paras = []
        if idx >= 0:
            _, ref_paras = extract_reference_paragraphs(doc)

        # 1. 写回格式化后的参考文献文本，并加书签
        for i, para in enumerate(ref_paras):
            if i >= len(formatted_refs):
                break
            # 清空段落
            for run in list(para.runs):
                run._r.getparent().remove(run._r)
            para.add_run(formatted_refs[i])
            add_bookmark_to_paragraph(para, f"Ref_{i + 1}")
            stats["bookmarks_added"] += 1

        # 2. 正文：先做重排文本替换（旧→新编号，两阶段标记交换避免级联覆盖）
        if order_map:
            from .utils.docx_utils import apply_two_phase_renumber
            body_paras = doc.paragraphs[:idx] if idx >= 0 else doc.paragraphs
            n = apply_two_phase_renumber(body_paras, order_map)
            stats["citations_renumbered"] += n

        # 3. 正文：把 [n] 替换为超链接（跳过参考文献区）
        body_paras = doc.paragraphs[:idx] if idx >= 0 else doc.paragraphs
        for para in body_paras:
            import re
            for m in re.finditer(r"\[(\d+)\]", para.text or ""):
                n = int(m.group(1))
                stats["hyperlinks_added"] += replace_citation_with_hyperlink(
                    para, n, superscript=superscript)

        doc.save(output_path)
        return stats

    def _merge_consecutive(self, para) -> int:
        """把段落内的连续引用 [a][b][c]... 合并为 [a-c]（仅纯文本 run 内相邻）。"""
        from .utils.docx_utils import _iter_direct_runs, _runs_text, _set_runs_text
        import re
        pattern = re.compile(r"(\[\d+\])((?:\[\d+\])+)")
        count = 0
        for run_el in list(_iter_direct_runs(para._p)):
            text = _runs_text(run_el)
            if not re.search(r"\]\[", text):
                continue
            def _repl(m):
                nums = [int(x) for x in re.findall(r"\[(\d+)\]", m.group(0))]
                if not nums:
                    return m.group(0)
                merged = []
                for n in sorted(nums):
                    if not merged:
                        merged.append(n)
                    else:
                        merged.append(n)
                # 压缩成范围 [1-3] 或单点
                groups = []
                start = prev = merged[0]
                for n in merged[1:]:
                    if n == prev + 1:
                        prev = n
                    else:
                        groups.append((start, prev))
                        start = prev = n
                groups.append((start, prev))
                parts = [f"{a}" if a == b else f"{a}-{b}" for a, b in groups]
                return "[" + ",".join(parts) + "]"
            new_text = pattern.sub(_repl, text)
            if new_text != text:
                _set_runs_text(run_el, new_text)
                count += 1
        return count
