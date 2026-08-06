"""Word (.docx) 底层操作：引用提取、书签、内部超链接、角标上标。

使用 python-docx 高层 API + 直接 OOXML（docx.oxml）操作：
  - 交叉引用用书签 + 内部 w:hyperlink（w:anchor 指向书签名），
    不需要修改 document.xml.rels（外部链接才需要）。
  - 超链接角标用 w:vertAlign val="superscript"。
  - 参考文献编号是字面文本（非 SEQ 域），重排时用两阶段标记交换避免级联覆盖。
"""

import re
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 参考文献标题关键词（按长度降序优先匹配）
REF_HEADINGS = [
    "参考文献", "References", "REFERENCES", "Reference", "REFERENCE",
    "Bibliography", "BIBLIOGRAPHY", "References（References）",
]

# 正文引用标记：方括号内的数字/逗号/连字符/空格
_CITE_RE = re.compile(r"\[([0-9,\-\s]+)\]")


def _parse_citation_range(content: str) -> list[int]:
    """解析 "[1,3-5]" → [1,3,4,5]。"""
    nums: list[int] = []
    for part in content.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, _, b = part.partition("-")
            try:
                start, end = int(a), int(b)
                if start <= end:
                    nums.extend(range(start, end + 1))
                else:
                    nums.extend(range(start, end - 1, -1))
            except ValueError:
                continue
        else:
            try:
                nums.append(int(part))
            except ValueError:
                continue
    return nums


def dedup_adjacent_citations(citations: list[int]) -> list[int]:
    """合并相邻重复，[1][1][2] → [1][2]。"""
    out = []
    for c in citations:
        if not out or out[-1] != c:
            out.append(c)
    return out


def find_references_heading_index(doc) -> int:
    """定位参考文献标题段落的下标；找不到返回 -1。"""
    paragraphs = doc.paragraphs
    # 从末尾向上找，最后一段是参考文献标题的可能性高
    for i in range(len(paragraphs) - 1, -1, -1):
        text = (paragraphs[i].text or "").strip()
        if not text:
            continue
        if text in REF_HEADINGS or re.match(r"^\s*[1-9]\d*\s*参考文献", text):
            return i
    # 宽松匹配：正文中含"参考文献"的短段
    for i, p in enumerate(paragraphs):
        t = (p.text or "").strip()
        if 1 <= len(t) <= 30 and ("参考文献" in t or "References" in t):
            return i
    return -1


def extract_reference_paragraphs(doc) -> tuple[int, list]:
    """返回 (start_index, paragraphs)：参考文献标题后到文末的段落列表。"""
    idx = find_references_heading_index(doc)
    paragraphs = doc.paragraphs
    if idx < 0:
        # 无参考文献标题：尝试从后向前找以 [N] 开头的段落
        start = len(paragraphs)
        for i in range(len(paragraphs) - 1, -1, -1):
            if re.match(r"^\s*\[\d+\]", paragraphs[i].text or ""):
                start = i
            else:
                break
        if start < len(paragraphs):
            return start, paragraphs[start:]
        return len(paragraphs), []
    refs = paragraphs[idx + 1:]
    # 去掉文末空段
    while refs and not (refs[-1].text or "").strip():
        refs.pop()
    return idx + 1, refs


def extract_citations_from_document(doc, ref_start_index: int) -> list[int]:
    """扫描正文（参考文献标题之前），返回按首次出现顺序去重后的引用编号。"""
    citations: list[int] = []
    paragraphs = doc.paragraphs
    upper = ref_start_index if ref_start_index is not None and ref_start_index >= 0 else len(paragraphs)
    for i, p in enumerate(paragraphs):
        if i >= upper:
            break
        text = p.text or ""
        if not text.strip():
            continue
        for m in _CITE_RE.finditer(text):
            content = m.group(1)
            if not re.search(r"\d", content):
                continue
            for n in _parse_citation_range(content):
                if n > 0 and n not in citations:
                    citations.append(n)
    return dedup_adjacent_citations(citations)


def _make_run(text: str, rpr_parent=None):
    """创建带可选 rPr 的 w:r。"""
    r = OxmlElement("w:r")
    if rpr_parent is not None:
        rpr = deepcopy(rpr_parent)
        if rpr is not None:
            r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def make_hyperlink_element(anchor: str, text: str, superscript: bool = False,
                           color: str | None = None):
    """构造内部交叉引用超链接元素（w:hyperlink + w:anchor=书签名）。

    返回可插入段落的 OxmlElement。superscript=True 时把角标设为上标。
    """
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), anchor)  # 内部锚点：指向书签名，而非外部关系
    hl.set(qn("w:history"), "1")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if superscript:
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rpr.append(va)
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rpr.append(c)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hl.append(r)
    return hl


def add_bookmark_to_paragraph(para, name: str) -> None:
    """给整个段落加书签（w:bookmarkStart 在首 run 前，w:bookmarkEnd 在段末）。"""
    p = para._p
    bid = abs(hash(name)) % 1000000
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    # bookmarkStart 插到段落第一个子元素之前
    first = p.find(qn("w:pPr"))
    insert_after = p.find(qn("w:r"))
    if insert_after is None:
        p.append(start)
    else:
        insert_after.addprevious(start)
    p.append(end)


def set_superscript(run) -> None:
    """把 python-docx Run 设为上标。"""
    try:
        run.font.superscript = True
    except Exception:
        rpr = run._r.get_or_add_rPr()
        va = rpr.find(qn("w:vertAlign"))
        if va is None:
            va = OxmlElement("w:vertAlign")
            rpr.append(va)
        va.set(qn("w:val"), "superscript")


def _iter_direct_runs(p):
    """遍历段落 p 的【直接】<w:r> 子元素（跳过超链接内部、书签等）。"""
    for child in p:
        if child.tag == qn("w:r"):
            yield child


def _runs_text(r):
    """读取 run 内全部 w:t 文本。"""
    parts = []
    for t in r.iter(qn("w:t")):
        parts.append(t.text or "")
    return "".join(parts)


def _set_runs_text(r, text: str):
    """清空 run 内 w:t 并写入新文本（保留 rPr）。"""
    for t in list(r.iter(qn("w:t"))):
        r.remove(t)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)


def replace_citation_with_hyperlink(para, ref_num: int, superscript: bool = True) -> int:
    """把段落内所有 `[ref_num]` 替换为内部超链接。返回替换次数。

    只处理段落的直接 <w:r> 子元素（避免嵌套超链接）。一个 run 内可能同时
    含多个 `[n]` 与前后文本；拆分会产生新 run，因此循环直到无匹配。
    """
    if ref_num <= 0:
        return 0
    target = f"[{ref_num}]"
    count = 0
    p = para._p
    while True:
        run_el = None
        for r in _iter_direct_runs(p):
            if target in _runs_text(r):
                run_el = r
                break
        if run_el is None:
            break
        text = _runs_text(run_el)
        rpr = run_el.find(qn("w:rPr"))
        # 把 text 按 target 切段，交替生成 普通run / 超链接run
        segments = text.split(target)
        new_nodes = []
        for i, seg in enumerate(segments):
            if seg:
                new_nodes.append(_make_run(seg, rpr))
            if i < len(segments) - 1:
                new_nodes.append(make_hyperlink_element(
                    f"Ref_{ref_num}", target, superscript=superscript))
                count += 1
        # 用新节点替换原 run（addprevious 一次只能插一个）
        for node in new_nodes:
            run_el.addprevious(node)
        p.remove(run_el)
    return count


def replace_citation_text_in_paragraph(para, old_num: int, new_num: int) -> int:
    """把段落内 `[old_num]` 文本替换为 `[new_num]`（纯文本，不含超链接）。"""
    old = f"[{old_num}]"
    new = f"[{new_num}]"
    count = 0
    for run_el in list(_iter_direct_runs(para._p)):
        text = _runs_text(run_el)
        if old in text:
            _set_runs_text(run_el, text.replace(old, new))
            count += text.count(old)
    return count


def apply_two_phase_renumber(paragraphs, old_to_new: dict) -> int:
    """两阶段重排：先 `[old]`→`__REORDER_<new>__`，再 `__REORDER_<new>__`→`[new]`。

    避免单次替换导致的级联覆盖（如 1→2 时旧 2 被误改）。
    """
    phase1 = {f"[{o}]": f"__REORDER_{n}__" for o, n in old_to_new.items() if o != n}
    phase2 = {f"__REORDER_{n}__": f"[{n}]" for o, n in old_to_new.items() if o != n}
    total = 0
    for para in paragraphs:
        for run_el in list(_iter_direct_runs(para._p)):
            text = _runs_text(run_el)
            new_text = text
            for a, b in phase1.items():
                new_text = new_text.replace(a, b)
            for a, b in phase2.items():
                new_text = new_text.replace(a, b)
            if new_text != text:
                _set_runs_text(run_el, new_text)
                total += 1
    return total
