"""格式转换：把规范化 ref dict 渲染成目标格式字符串。"""

from .formats.gbt7714 import GBT7714Formatter
from .formats.ieee import IEEEFormatter
from .formats.apa7 import APA7Formatter
from .formats.chicago import ChicagoFormatter
from .formats.mla import MLAFormatter
from .formats.harvard import HarvardFormatter


class ReferenceFormatter:
    """6 种标准格式转换。"""

    FORMATS = {
        "gbt7714": GBT7714Formatter,
        "ieee": IEEEFormatter,
        "apa7": APA7Formatter,
        "chicago": ChicagoFormatter,
        "mla": MLAFormatter,
        "harvard": HarvardFormatter,
    }

    LABELS = {
        "gbt7714": "GB/T 7714-2015",
        "ieee": "IEEE",
        "apa7": "APA 7th",
        "chicago": "Chicago 17th",
        "mla": "MLA 9th",
        "harvard": "Harvard",
    }

    def get_formatter(self, format_key: str):
        key = format_key or "gbt7714"
        cls = self.FORMATS.get(key, GBT7714Formatter)
        return cls()

    def format(self, references: list[dict], format_key: str):
        """把参考文献列表渲染为 `[N] 格式化文本`。返回 (列表, 修正数)。"""
        fmt = self.get_formatter(format_key)
        fixed = 0
        formatted = []
        for i, ref in enumerate(references, start=1):
            text = fmt.format(ref).strip()
            if text and ref.get("raw") and text not in ref.get("raw", ""):
                fixed += 1
            formatted.append(f"[{i}] {text}" if text else f"[{i}]")
        return formatted, fixed

    def format_single(self, ref: dict, format_key: str, number: int = None) -> str:
        fmt = self.get_formatter(format_key)
        text = fmt.format(ref).strip()
        if number:
            return f"[{number}] {text}"
        return text

    def write_back(self, input_path: str, output_path: str, formatted_refs: list[str],
                   merge: bool = False) -> None:
        """纯文本写回：替换参考文献区段落（不加交叉引用）。

        merge=True 时把正文中连续引用 [1][2] 合并为 [1-2]。
        """
        from docx import Document
        doc = Document(input_path)
        from .utils.docx_utils import (
            extract_reference_paragraphs, find_references_heading_index,
            _iter_direct_runs, _runs_text, _set_runs_text)
        idx = find_references_heading_index(doc)
        _, ref_paras = extract_reference_paragraphs(doc)
        for i, para in enumerate(ref_paras):
            if i >= len(formatted_refs):
                break
            # 清空段落文本
            for run in list(para.runs):
                run._r.getparent().remove(run._r)
            run = para.add_run(formatted_refs[i])

        if merge:
            import re
            body_paras = doc.paragraphs[:idx] if idx >= 0 else doc.paragraphs
            pattern = re.compile(r"(\[\d+\])((?:\[\d+\])+)")
            for para in body_paras:
                for run_el in list(_iter_direct_runs(para._p)):
                    text = _runs_text(run_el)
                    if not re.search(r"\]\[", text):
                        continue
                    def _repl(m):
                        nums = sorted(int(x) for x in re.findall(r"\[(\d+)\]", m.group(0)))
                        groups = []
                        start = prev = nums[0]
                        for n in nums[1:]:
                            if n == prev + 1:
                                prev = n
                            else:
                                groups.append((start, prev)); start = prev = n
                        groups.append((start, prev))
                        parts = [f"{a}" if a == b else f"{a}-{b}" for a, b in groups]
                        return "[" + ",".join(parts) + "]"
                    new_text = pattern.sub(_repl, text)
                    if new_text != text:
                        _set_runs_text(run_el, new_text)

        doc.save(output_path)
