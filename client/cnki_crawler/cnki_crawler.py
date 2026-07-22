#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
鸿讯 HONGXUN · CNKI 知网文献抓取工具 v2.0
软件著作权登记版
基于 Selenium + 用户已登录 Chrome 调试模式

功能：自动化提取论文标题、摘要、作者、作者单位、期刊名称、发表时间
输出：与 engine.py 格式一致的 DOCX 报告
"""
import os, re, time, random, json, signal, sys, subprocess, shutil, zipfile, io, textwrap, ssl
import urllib.request
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# selenium 按需导入（在 connect_chrome 时引入）
# from selenium import webdriver
# from selenium.webdriver.common.by import By
# from selenium.webdriver.support.ui import WebDriverWait
# from selenium.webdriver.support import expected_conditions as EC

# ============ 全局中断处理 ============
_sigint_count = 0

def _signal_handler(sig, frame):
    global _sigint_count
    _sigint_count += 1
    print("\n\n⚠️ 收到中断信号（Ctrl+C），正在终止...")
    if _sigint_count >= 2:
        print("   强制退出。")
        sys.exit(1)

signal.signal(signal.SIGINT, _signal_handler)

# ============ 配置 ============
CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".cnki_crawler_config.json")
PAGE_DELAY = (4, 7)
DETAIL_DELAY = (3, 6)
MAX_CLICK_RETRY = 3

# 中国可访问的资源镜像
CHROMEDRIVER_MIRRORS = [
    "https://registry.npmmirror.com/-/binary/chrome-for-testing",
    "https://storage.googleapis.com/chrome-for-testing-public",
]
CHROME_VERSION_LOOKUP = "https://googlechromelabs.github.io/chrome-for-testing"
CHROME_DMG_MIRRORS = [
    "https://dl.google.com/chrome/mac/universal/stable/GGRO/googlechrome.dmg",
    "https://dldir1.qq.com/google/chrome/googlechrome.dmg",
]

# ============ 环境检查与自动安装 ============
CHROME_PATHS_MAC = [
    "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
    os.path.expanduser("~/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"),
]
CHROME_PATHS_WIN = [
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    os.path.expanduser(r"~\AppData\Local\Google\Chrome\Application\chrome.exe"),
]


def _get_platform():
    if sys.platform == "darwin":
        return "mac"
    elif sys.platform == "win32":
        return "win"
    return "linux"


def _find_chrome() -> str | None:
    if _get_platform() == "mac":
        for p in CHROME_PATHS_MAC:
            if os.path.exists(p):
                return p
        try:
            r = subprocess.run(["brew", "--prefix", "--cask", "google-chrome"],
                               capture_output=True, text=True, timeout=10)
            if r.returncode == 0:
                bp = r.stdout.strip()
                candidate = os.path.join(bp, "Google Chrome.app/Contents/MacOS/Google Chrome")
                if os.path.exists(candidate):
                    return candidate
        except: pass
    elif _get_platform() == "win":
        for p in CHROME_PATHS_WIN:
            if os.path.exists(p):
                return p
    return None


def _get_chrome_version(chrome_path: str) -> str | None:
    try:
        r = subprocess.run([chrome_path, "--version"], capture_output=True, text=True, timeout=10)
        ver_str = re.search(r'(\d+\.\d+\.\d+\.\d+)', r.stdout or r.stderr)
        if ver_str:
            return ver_str.group(1)
    except: pass
    return None


def _get_chromedriver_installed_version() -> str | None:
    paths_to_try = ["chromedriver"]
    if sys.platform == "darwin":
        local_cd = os.path.expanduser("~/.local/bin/chromedriver")
        if os.path.exists(local_cd):
            paths_to_try.insert(0, local_cd)
    for cd_path in paths_to_try:
        try:
            r = subprocess.run([cd_path, "--version"], capture_output=True, text=True, timeout=10)
            ver_str = re.search(r'(\d+\.\d+\.\d+\.\d+)', r.stdout or r.stderr)
            if ver_str:
                return ver_str.group(1)
        except: pass
    return None


def _download_url(url: str, timeout: int = 120) -> bytes | None:
    """下载 URL 内容，带重试"""
    _ssl_ctx = ssl._create_unverified_context()
    for attempt in range(3):
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=timeout, context=_ssl_ctx) as resp:
                return resp.read()
        except Exception:
            time.sleep(2)
            continue
    return None


def _install_chromedriver(chrome_version: str) -> bool:
    """自动下载并安装匹配的 ChromeDriver（优先使用中国大陆可访问的 npmmirror）"""
    major = chrome_version.split(".")[0]
    print(f"  正在查找 ChromeDriver version {major}...")

    platform_map = {"mac": "mac-x64", "win": "win32", "linux": "linux64"}
    plat = platform_map.get(_get_platform(), "mac-x64")
    if _get_platform() == "mac":
        import platform
        if platform.machine() == "arm64":
            plat = "mac-arm64"

    # 查最新版本号
    latest_ver = None
    print(f"  正在查询 ChromeDriver（for Chrome {major}）...")
    for retry in range(3):
        try:
            ver_url = f"{CHROME_VERSION_LOOKUP}/LATEST_RELEASE_{major}"
            data = _download_url(ver_url, timeout=15)
            if data:
                latest_ver = data.decode().strip()
                break
        except Exception:
            time.sleep(1)

    # 构造下载 URL（npmmirror 优先）
    dl_urls = []
    if latest_ver:
        dl_urls.append(f"{CHROMEDRIVER_MIRRORS[0]}/{latest_ver}/{plat}/chromedriver-{plat}.zip")
        dl_urls.append(f"{CHROMEDRIVER_MIRRORS[1]}/{latest_ver}/{plat}/chromedriver-{plat}.zip")
    # fallback
    dl_urls.append(f"{CHROMEDRIVER_MIRRORS[1]}/{major}.0.0.0/{plat}/chromedriver-{plat}.zip")

    zip_data = None
    for dl_url in dl_urls:
        print(f"  正在下载 ChromeDriver...")
        zip_data = _download_url(dl_url, timeout=120)
        if zip_data:
            break

    if zip_data is None:
        print(f"  ❌ 下载失败，请检查网络连接")
        return False

    # 解压安装
    try:
        exe_name = "chromedriver.exe" if _get_platform() == "win" else "chromedriver"
        with zipfile.ZipFile(io.BytesIO(zip_data)) as zf:
            for name in zf.namelist():
                if os.path.basename(name) == exe_name:
                    zf.extract(name, "/tmp/chromedriver_install")
                    src = os.path.join("/tmp/chromedriver_install", name)
                    dst_dir = os.path.expanduser("~/.local/bin") if _get_platform() == "mac" else "/usr/local/bin"
                    os.makedirs(dst_dir, exist_ok=True)
                    dst = os.path.join(dst_dir, exe_name)
                    shutil.copy2(src, dst)
                    os.chmod(dst, 0o755)
                    if sys.platform == "darwin":
                        try:
                            subprocess.run(["xattr", "-d", "com.apple.quarantine", dst],
                                           capture_output=True, timeout=10)
                        except Exception:
                            pass
                    shutil.rmtree("/tmp/chromedriver_install", ignore_errors=True)
                    print(f"  ✅ ChromeDriver 已安装到: {dst}")
                    # 加入 PATH
                    os.environ["PATH"] = f"{dst_dir}:{os.environ.get('PATH', '')}"
                    return True
    except Exception as e:
        print(f"  ❌ 解压/安装失败: {e}")
        return False
    return False


# ============ DOCX 报告生成（与 engine.py 格式一致）============

def _set_cell_shading(cell, color_hex: str):
    from docx.oxml.ns import qn
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def _style_cell_key(cell, color='F8FAFC'):
    _set_cell_shading(cell, color)
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True


def _set_run_font(run, text: str, size=Pt(10.5), bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.bold = bold
    from docx.oxml.ns import qn
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'SimSun')


def _set_paragraph_spacing(paragraph, line_spacing=Pt(18)):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _init_docx_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = Pt(18)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Pt(0)
    from docx.oxml.ns import qn
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = style.element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'SimSun')


def build_cnki_docx(papers: list[dict], file_path: str):
    """
    构造 CNKI 论文报告的 DOCX 文件。
    格式与 engine.py 的 _build_docx 一致：
      - 西文 Times New Roman，中文 宋体，10.5pt
      - 行间距固定 18 磅
      - 元信息表 + 逐篇论文
    papers 每项字段: title, authors, affiliations, journal, pub_date, abstract
    """
    doc = Document()
    _init_docx_styles(doc)

    # 标题
    title = doc.add_heading(f'CNKI 知网论文检索报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息表
    meta = doc.add_table(rows=3, cols=2, style='Table Grid')
    meta_data = [
        ("导出时间", datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ("来源", "中国知网 (CNKI)"),
        ("命中总数", f"{len(papers)} 篇"),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
    _style_cell_key(meta.cell(0, 0), 'EFF6FF')
    _style_cell_key(meta.cell(1, 0), 'EFF6FF')
    _style_cell_key(meta.cell(2, 0), 'EFF6FF')

    doc.add_paragraph()

    # 逐篇论文
    for idx, p in enumerate(papers, 1):
        doc.add_heading(f'{idx}. {p.get("title", "")}', level=2)

        # 字段行
        field_lines = [
            ("作者", p.get('authors', '')),
            ("作者单位", p.get('affiliations', '')),
            ("期刊名称", p.get('journal', '')),
            ("发表时间", p.get('pub_date', '')),
        ]
        for label, value in field_lines:
            if not value:
                continue
            line = doc.add_paragraph(f'【{label}】{value}')
            _set_paragraph_spacing(line)

        abstract = p.get('abstract', '')
        if abstract:
            ab_p = doc.add_paragraph(f'摘要：{abstract}')
            ab_p.paragraph_format.first_line_indent = Pt(21)
        else:
            doc.add_paragraph('摘要：无摘要')

    doc.save(file_path)
    return file_path


# ============ Selenium 连接 ============

# 在模块级别查找 chromedriver，避免每次连接时重复搜索
_CHROMEDRIVER_PATH = None
def _resolve_chromedriver():
    global _CHROMEDRIVER_PATH
    if _CHROMEDRIVER_PATH:
        return _CHROMEDRIVER_PATH
    candidates = ["chromedriver"]
    if sys.platform == "darwin":
        local_cd = os.path.expanduser("~/.local/bin/chromedriver")
        if os.path.exists(local_cd):
            candidates.insert(0, local_cd)
    for path in candidates:
        try:
            r = subprocess.run([path, "--version"], capture_output=True, text=True, timeout=5)
            if r.returncode == 0:
                _CHROMEDRIVER_PATH = path
                return path
        except Exception:
            continue
    return "chromedriver"  # fallback to PATH


def connect_chrome(timeout: int = 15):
    """
    连接到已开启的 Chrome 调试端口（127.0.0.1:9222）。

    关键解决：
    - 首次连接时杀死残留 chromedriver 进程，避免端口冲突
    - 进程环境严格禁用 Selenium Manager（国内网络被墙卡死的主因）
    - 使用显式 chromedriver 路径
    - 线程超时保护兜底
    """
    import os as _os
    import subprocess as _sp

    # 1) 在进程级和环境变量双重禁用 Selenium Manager 联网
    #    （Selenium Manager 会在创建 ChromeService 时检查网络）
    _os.environ['SE_MANAGER_ENABLED'] = 'false'
    _os.environ['SE_MANAGER_DOWNLOAD_CHECK'] = 'false'

    # 2) 杀死任何残留的 chromedriver 进程，确保端口 9515 可用
    _sp.run(['pkill', '-9', '-f', 'chromedriver'], capture_output=True, timeout=5)

    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.chrome.service import Service
    import threading

    cd_path = _resolve_chromedriver()

    opts = Options()
    opts.add_experimental_option("debuggerAddress", "127.0.0.1:9222")

    service = Service(executable_path=cd_path)

    result = [None]
    exception = [None]
    done = threading.Event()

    def _connect():
        try:
            driver = webdriver.Chrome(service=service, options=opts)
            driver.set_page_load_timeout(30)
            result[0] = driver
        except Exception as e:
            exception[0] = e
        finally:
            done.set()

    t = threading.Thread(target=_connect, daemon=True)
    t.start()

    if not done.wait(timeout=timeout):
        # 超时后清理 chromedriver 残骸
        _sp.run(['pkill', '-9', '-f', 'chromedriver'], capture_output=True, timeout=5)
        raise TimeoutError(
            f"Chrome 连接超时（{timeout}秒）" +
            "\n常见原因：国内网络下 Selenium Manager 联网被墙，请重试" +
            "\n请确认 Chrome 已以调试模式启动（端口 9222）"
        )
    if exception[0]:
        raise exception[0]

    return result[0]


def safe_click(driver, elem, retries=MAX_CLICK_RETRY):
    for _ in range(retries):
        try:
            driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elem)
            time.sleep(0.2)
            driver.execute_script("arguments[0].click();", elem)
            return True
        except Exception:
            time.sleep(0.8)
    return False


def turn_next_page(driver) -> bool:
    """翻到下一页，返回 False 表示没有下一页"""
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC

    selectors = [
        ("css", "a.next"),
        ("css", "a.page-next"),
        ("css", "li.layui-laypage-next"),
        ("css", "a[title='下一页']"),
        ("xpath", "//a[contains(text(),'下一页')]"),
        ("xpath", "//a[contains(@class,'next')]"),
        ("xpath", "//*[contains(@aria-label,'下一页')]"),
    ]
    for typ, sel in selectors:
        try:
            elem = WebDriverWait(driver, 3).until(
                EC.element_to_be_clickable((By.XPATH if typ == "xpath" else By.CSS_SELECTOR, sel))
            )
            # 检查是否 disabled
            classes = elem.get_attribute("class") or ""
            if "disabled" in classes:
                return False
            if safe_click(driver, elem):
                WebDriverWait(driver, 20).until(
                    EC.presence_of_element_located(
                        (By.XPATH, "//a[contains(@href,'article') or contains(@href,'detail') or contains(@href,'doi') or contains(@href,'kns')]")
                    )
                )
                time.sleep(random.uniform(*PAGE_DELAY))
                return True
        except Exception:
            continue
    return False


def find_all_paper_links(driver):
    """查找论文标题和可点击的链接，返回 [(标题, 链接元素), ...]"""
    from selenium.webdriver.common.by import By
    seen_href = set()
    seen_title = set()
    results = []

    # URL 模式匹配
    patterns = [
        "//a[contains(@href,'kcms2/article/')]",
        "//a[contains(@href,'kcms/detail/')]",
        "//a[contains(@href,'kns8/detail/')]",
        "//a[contains(@href,'kns8s/detail/')]",
        "//a[contains(@href,'link.cnki.net/doi/')]",
        "//a[contains(@href,'article/')]",
    ]
    all_links = []
    for pat in patterns:
        for link in driver.find_elements(By.XPATH, pat):
            try:
                href = link.get_attribute("href")
                if href and href not in seen_href and link.is_displayed():
                    all_links.append((link, href))
                    seen_href.add(href)
            except: continue

    # 从附近元素获取真实标题
    for link, href in all_links:
        text = link.text.strip()
        for xpath in [
            "./ancestor::tr[1]//a[contains(@class,'title')]",
            "./ancestor::tr[1]//*[contains(@class,'title')]",
            "./ancestor::div[contains(@class,'result')]//*[contains(@class,'title')]",
            "./ancestor::div[contains(@class,'list-item')]//*[contains(@class,'title')]",
            "./preceding::*[contains(@class,'title')][1]",
            "./ancestor::tr[1]//td[1]",
        ]:
            try:
                t = link.find_element(By.XPATH, xpath).text.strip()
                if t and len(t) > 6 and "总库" not in t:
                    text = t; break
            except: continue
        if text and len(text) > 5 and text not in seen_title:
            results.append((text, link))
            seen_title.add(text)

    if results:
        return results

    # CSS 选择器回退
    css_sels = [
        "p.title a", "a.name", ".title a", "p.title", ".title",
        "[class*='title'] a", "[class*='name'] a",
    ]
    for sel in css_sels:
        for el in driver.find_elements(By.CSS_SELECTOR, sel):
            try:
                text = el.text.strip()
                href = el.get_attribute("href") or ""
                if text and len(text) > 6 and "总库" not in text and text not in seen_title:
                    results.append((text, el))
                    seen_title.add(text)
            except: continue
        if results:
            break
    return results


def _extract_from_detail_page(driver) -> dict:
    """
    从知网论文详情页提取完整信息。
    返回: {
        "title": str,
        "abstract": str,
        "authors": str,
        "affiliations": str,
        "journal": str,
        "pub_date": str,
    }
    """
    from selenium.webdriver.common.by import By

    result = {"title": "", "abstract": "", "authors": "",
              "affiliations": "", "journal": "", "pub_date": ""}

    # ---- 展开摘要 ----
    more_selectors = [
        ("css", "a.more, span.more, a[class*='more']"),
        ("xpath", "//a[contains(text(),'更多')]"),
        ("xpath", "//a[contains(text(),'展开')]"),
        ("xpath", "//span[contains(text(),'更多')]"),
        ("xpath", "//span[contains(text(),'展开')]"),
    ]
    for typ, sel in more_selectors:
        try:
            btn = driver.find_element(By.XPATH if typ == "xpath" else By.CSS_SELECTOR, sel)
            if btn.is_displayed():
                safe_click(driver, btn)
                time.sleep(0.5)
                break
        except: continue

    # 下拉页面
    try:
        driver.execute_script("window.scrollTo(0, 600);")
    except: pass
    time.sleep(1)

    # ---- 标题 ----
    try:
        t = driver.title
        for suffix in ["_知网", "_中国知网", "- CNKI", " - 中国知网", "_CNKI", "知网"]:
            t = t.replace(suffix, "").strip().rstrip("-").strip()
        if t and len(t) > 6 and "总库" not in t:
            result["title"] = t
    except: pass
    if not result["title"]:
        for sel in ["meta[name='citation_title']", "h1", "h2", "h1.title", "h2.title",
                     ".detail-title", ".paper-title", ".article-title", ".title", "[class*='title']"]:
            try:
                el = driver.find_element(By.CSS_SELECTOR, sel)
                if el.tag_name == "meta":
                    result["title"] = el.get_attribute("content").strip()
                else:
                    result["title"] = el.text.strip()
                if result["title"] and len(result["title"]) > 6 and "总库" not in result["title"]:
                    break
                result["title"] = ""
            except: continue

    # ---- 摘要 ----
    for sel in ["#ChDivSummary", "#abstractContent", ".abstract-content",
                 "div.abstract", "p.abstract", "div.abstract-text",
                 "[class*='abstract']"]:
        try:
            text = driver.find_element(By.CSS_SELECTOR, sel).text.strip()
            if text and len(text) > 30:
                result["abstract"] = text
                break
        except: continue

    # ---- 作者 ----
    for sel in ["div.author", "span.author", "p.author", "a.author",
                 ".author-info", "[class*='author']"]:
        try:
            text = driver.find_element(By.CSS_SELECTOR, sel).text.strip()
            if text:
                text = re.sub(r'\[?\d+(?:[,，]\d*)*\]?', '', text)
                text = re.sub(r'[¹²³⁰-⁹]', '', text)
                result["authors"] = text
                break
        except: continue

    # ---- 作者单位 ----
    # CNKI 详情页作者单位通常在 #authorAffix 或 .affiliation 或 .author-affiliation
    for sel in ["#authorAffix", ".affiliation", ".author-affiliation",
                 "[class*='affiliation']", ".orgn", "[class*='orgn']",
                 "p.affiliation", "div.affiliation"]:
        try:
            el = driver.find_element(By.CSS_SELECTOR, sel)
            text = el.text.strip()
            # 去掉数字标记
            text = re.sub(r'^\d+[.．、\s]*', '', text)
            if text:
                result["affiliations"] = text
                break
        except: continue
    # 备选：从页面 body 文本中提取机构信息
    if not result["affiliations"]:
        try:
            body = driver.find_element(By.TAG_NAME, "body").text
            # 常见模式："作者单位" 或 "机构" 后跟的内容
            for kw in ["作者单位", "机构：", "机构:", "作者机构"]:
                idx = body.find(kw)
                if idx >= 0:
                    snippet = body[idx:idx+300].split("\n")[0].split(".")[0][:200]
                    result["affiliations"] = snippet.replace(kw, "").strip()
                    if result["affiliations"]:
                        break
        except: pass

    # ---- 期刊 ----
    for sel in ["div.source", "span.source", "p.source",
                 "meta[name='citation_journal_title']",
                 ".journal-title", "[class*='source']", "[class*='jour']"]:
        try:
            el = driver.find_element(By.CSS_SELECTOR, sel)
            if el.tag_name == "meta":
                result["journal"] = el.get_attribute("content").strip()
            else:
                result["journal"] = el.text.strip()
            if result["journal"]:
                break
        except: continue

    # ---- 发表时间 ----
    for sel in ["meta[name='citation_date']", "meta[name='citation_publication_date']",
                 "meta[name='citation_year']", ".year", "[class*='year']",
                 ".date", "[class*='date']"]:
        try:
            el = driver.find_element(By.CSS_SELECTOR, sel)
            if el.tag_name == "meta":
                content = el.get_attribute("content").strip()
            else:
                content = el.text.strip()
            yr = re.search(r'\b(19|20)\d{2}\b', content)
            if yr:
                result["pub_date"] = yr.group(0)
                break
        except: continue
    if not result["pub_date"]:
        try:
            body_text = driver.find_element(By.TAG_NAME, "body").text[:3000]
            yr = re.search(r'\b(19|20)\d{2}\b', body_text)
            if yr:
                result["pub_date"] = yr.group(0)
        except: pass

    return result


# ============ 环境检查 ============

def check_environment() -> dict:
    """
    检查 Chrome 和 ChromeDriver 环境。
    返回: {"ok": bool, "chrome": bool, "chrome_version": str|None,
           "chromedriver": bool, "message": str}
    """
    import subprocess as _sp
    result = {"ok": False, "chrome": False, "chrome_version": None,
              "chromedriver": False, "message": ""}

    chrome_path = _find_chrome()
    if not chrome_path:
        result["message"] = "未检测到 Google Chrome 浏览器"
        return result
    result["chrome"] = True

    chrome_ver = _get_chrome_version(chrome_path)
    result["chrome_version"] = chrome_ver

    # 检查 chromedriver
    driver_ver = _get_chromedriver_installed_version()
    if driver_ver:
        result["chromedriver"] = True

    if result["chrome"] and result["chromedriver"]:
        result["ok"] = True
        result["message"] = "环境就绪"
    else:
        result["message"] = f"Chrome {'✓' if result['chrome'] else '✗'} | ChromeDriver {'✓' if result['chromedriver'] else '✗'}"

    return result


def auto_fix_environment(progress_callback=None) -> bool:
    """
    自动修复环境（安装 ChromeDriver），返回 True 表示修复成功。
    """
    def _log(msg):
        if progress_callback:
            progress_callback(msg)
        else:
            print(msg)

    _log("正在检测环境...")
    env = check_environment()

    if env["ok"]:
        _log("✅ 环境就绪")
        return True

    if not env["chrome"]:
        _log("❌ 未检测到 Google Chrome")
        _log("请先安装 Google Chrome 浏览器（https://www.google.com/chrome/）")
        # macOS 尝试自动安装
        if sys.platform == "darwin":
            _log("正在通过 Homebrew 安装 Chrome...")
            try:
                r = subprocess.run(["brew", "install", "--cask", "google-chrome"],
                                   capture_output=True, text=True, timeout=300)
                if r.returncode == 0:
                    _log("✅ Chrome 安装完成")
                else:
                    _log("Homebrew 安装失败，请手动安装 Chrome 后重试")
                    return False
            except Exception:
                # 下载 dmg
                _log("正在下载 Chrome...")
                dmg_data = None
                for url in CHROME_DMG_MIRRORS:
                    dmg_data = _download_url(url, timeout=300)
                    if dmg_data:
                        break
                if not dmg_data:
                    _log("❌ Chrome 下载失败，请手动安装")
                    return False
                _log("正在安装 Chrome...")
                dmg_path = "/tmp/googlechrome.dmg"
                with open(dmg_path, "wb") as f:
                    f.write(dmg_data)
                try:
                    subprocess.run(["hdiutil", "attach", dmg_path, "-quiet", "-nobrowse"],
                                   capture_output=True, timeout=30)
                    subprocess.run(["cp", "-r", "/Volumes/Google Chrome/Google Chrome.app",
                                    "/Applications/"],
                                   capture_output=True, timeout=60)
                    subprocess.run(["hdiutil", "detach", "/Volumes/Google Chrome", "-quiet"],
                                   capture_output=True, timeout=30)
                    os.remove(dmg_path)
                    _log("✅ Chrome 安装完成")
                except Exception as e:
                    _log(f"❌ Chrome 安装失败: {e}")
                    return False
        else:
            return False

    # 安装/更新 ChromeDriver
    chrome_ver = env["chrome_version"] or _get_chrome_version(_find_chrome())
    if not chrome_ver:
        _log("❌ 无法确定 Chrome 版本")
        return False

    _log(f"正在安装 ChromeDriver（匹配 Chrome {chrome_ver.split('.')[0]}）...")
    success = _install_chromedriver(chrome_ver)
    if success:
        _log("✅ ChromeDriver 安装完成")
        return True
    else:
        _log("❌ ChromeDriver 安装失败")
        return False


# ============ 浏览器调试模式启动 ============

# 鸿讯独立 Chrome 配置目录（与用户日常 Chrome 隔离，知网登录仅需一次）
HONGXUN_CHROME_DIR = os.path.join(os.path.expanduser("~"), ".hongxun", "chrome-profile")


def launch_chrome_debug() -> bool:
    """
    以调试模式启动 Chrome（端口 9222），返回是否成功。

    使用独立的用户数据目录（~/.hongxun/chrome-profile/），
    确保：
    - 不影响用户已有的 Chrome 会话
    - 调试端口可靠绑定（解决了 open -na 在 Chrome 已运行时失效的问题）
    - 知网登录状态持久保存，后续使用无需重新登录
    """
    import socket

    # 先检查端口 9222 是否已就绪
    try:
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(1)
        result = sock.connect_ex(('127.0.0.1', 9222))
        sock.close()
        if result == 0:
            return True
    except Exception:
        pass

    chrome_path = _find_chrome()
    if not chrome_path:
        return False

    # 确保独立配置目录存在
    os.makedirs(HONGXUN_CHROME_DIR, exist_ok=True)

    # 清理上次崩溃遗留的锁文件，确保新实例正常启动
    for lockfile in ["SingletonLock", "SingletonCookie", "SingletonSocket"]:
        lock_path = os.path.join(HONGXUN_CHROME_DIR, lockfile)
        if os.path.exists(lock_path):
            try:
                os.remove(lock_path)
            except Exception:
                pass

    try:
        # 直接用二进制路径启动（不用 open -na），配合独立配置目录，
        # 即使 Chrome 已在运行也能可靠绑定调试端口
        subprocess.Popen(
            [
                chrome_path,
                "--remote-debugging-port=9222",
                f"--user-data-dir={HONGXUN_CHROME_DIR}",
                "--no-first-run",
                "--no-default-browser-check",
                "--disable-default-apps",
                "--disable-sync",
                "--disable-features=TranslateUI,ChromeWhatsNewUI",
                "https://kns.cnki.net",
            ],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )

        # 等待端口 9222 就绪（最长 20 秒）
        for _ in range(40):
            time.sleep(0.5)
            try:
                sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                sock.settimeout(1)
                result = sock.connect_ex(('127.0.0.1', 9222))
                sock.close()
                if result == 0:
                    return True
            except Exception:
                pass
        return False
    except Exception:
        return False


# ============ 一键抓取流程 ============

def crawl_all_pages(driver, progress_callback=None) -> list[dict]:
    """
    在已连接的 Chrome 上执行知网检索结果页的逐页抓取。
    - 从当前页面开始，查找论文链接
    - 逐个打开详情页提取信息
    - 只关闭详情页标签，检索页始终保持
    - 翻页直到没有下一页
    返回 [paper_dict, ...]
    """
    from selenium.webdriver.common.by import By

    papers = []
    exist_titles = set()
    current_page = 1

    def _progress(msg, cur=0, total=0, pct=0):
        if progress_callback:
            progress_callback(msg, cur, total, pct)

    while True:
        _progress(f"正在处理第 {current_page} 页...", len(papers), 0,
                  0.3 + min(0.6, len(papers) * 0.01))

        paper_links = find_all_paper_links(driver)

        if not paper_links:
            try:
                snippet = driver.find_element(By.TAG_NAME, 'body').text[:200]
            except Exception:
                snippet = ""
            _progress(f"未找到论文链接（页面内容：{snippet[:60]}）", len(papers), 0, 1.0)
            break

        total_on_page = len(paper_links)
        _progress(f"本页共 {total_on_page} 篇论文", len(papers), total_on_page,
                  0.3 + min(0.6, len(papers) * 0.01))

        for idx, (title_text, link) in enumerate(paper_links):
            if title_text in exist_titles:
                continue

            _progress(f"正在获取 ({idx+1}/{total_on_page}) {title_text[:30]}...",
                      len(papers), total_on_page,
                      0.3 + min(0.6, (len(papers) + idx / max(1, total_on_page)) * 0.01))

            # 从结果行提取基本信息
            author = ""
            journal = ""
            try:
                parent_row = link.find_element(By.XPATH, "./ancestor::tr[1]")
                ael = parent_row.find_elements(By.CSS_SELECTOR, "td[class*='author'],[class*='author']")
                if ael:
                    author = ael[0].text.strip().split(";")[0].split(",")[0].strip()
                jel = parent_row.find_elements(By.CSS_SELECTOR, "td[class*='source'],[class*='source']")
                if jel:
                    journal = jel[0].text.strip()
            except Exception:
                pass

            orig_handle = driver.current_window_handle
            href = link.get_attribute("href")
            if not href:
                continue

            # 开新标签打开详情页（不关闭检索页）
            new_handle = None
            try:
                driver.switch_to.new_window('tab')
                new_handle = driver.current_window_handle
            except Exception:
                driver.execute_script("window.open('');")
                time.sleep(0.5)
                for h in driver.window_handles:
                    if h != orig_handle:
                        new_handle = h
                        break
            if not new_handle:
                continue

            driver.switch_to.window(new_handle)
            try:
                driver.get(href)
                time.sleep(random.uniform(2, 3))

                # 提取详情信息
                detail = _extract_from_detail_page(driver)
                final_title = detail["title"] or title_text
                final_author = detail["authors"] or author
                paper_entry = {
                    "title": final_title,
                    "authors": final_author,
                    "affiliations": detail.get("affiliations", ""),
                    "journal": detail.get("journal", "") or journal,
                    "pub_date": detail.get("pub_date", ""),
                    "abstract": detail.get("abstract", ""),
                }
                papers.append(paper_entry)
                exist_titles.add(title_text)

                _progress(f"✓ {final_title[:30]} | {final_author[:15]} | "
                          f"{paper_entry['journal'][:20]}",
                          len(papers), total_on_page,
                          0.3 + min(0.6, (len(papers)) * 0.01))

            except Exception as e:
                _progress(f"⚠ 提取失败: {str(e)[:40]}", len(papers), total_on_page, 0)
            finally:
                # 关闭详情页，回到检索页
                try:
                    driver.close()
                except Exception:
                    pass
                try:
                    if orig_handle in driver.window_handles:
                        driver.switch_to.window(orig_handle)
                    else:
                        driver.switch_to.window(driver.window_handles[0])
                except Exception:
                    pass

            time.sleep(random.uniform(0.5, 1.0))

        # 翻页
        if turn_next_page(driver):
            current_page += 1
        else:
            _progress(f"✅ 全部完成，共获取 {len(papers)} 篇论文", len(papers), 0, 1.0)
            break

    return papers


# ============ CLI 版帮助 ============

USAGE_TEXT = """
📖 CNKI 知网文献抓取工具 v2.0 — 使用说明

━━━ 功能概述 ━━━

从中国知网（CNKI）检索结果页中批量提取论文信息：
  · 论文标题
  · 作者姓名
  · 作者单位
  · 期刊名称
  · 发表时间
  · 论文摘要
保存为格式规范的 DOCX 报告文件。

━━━ 使用流程 ━━━

1. 环境检查
   自动检测 Chrome 和 ChromeDriver，缺失时自动安装。

2. 打开 Chrome
   系统自动以调试模式（端口 9222）启动 Chrome。

3. 登录知网并检索
   在打开的 Chrome 中登录知网（kns.cnki.net），
   进行检索操作，进入结果页。

4. 自动抓取
   点击「我已登录并完成检索」后自动逐页抓取所有论文。

━━━ CLI 命令 ━━━

  python3 cnki_crawler.py         交互式运行
  python3 cnki_crawler.py --help   显示此帮助
"""


def print_usage():
    print(textwrap.dedent(USAGE_TEXT).strip())


def cli_main():
    """CLI 交互入口（非 GUI 模式时使用）"""
    import socket

    print("=" * 60)
    print("  知网文献抓取工具 v2.0")
    print("=" * 60)

    # 环境检查
    if not check_environment()["ok"]:
        print("\n⚠ 环境未就绪，尝试自动修复...")
        if not auto_fix_environment():
            print("❌ 环境修复失败，请手动安装 Chrome 和 ChromeDriver")
            return
    else:
        print("✅ 环境就绪")

    # 保存路径
    default_path = os.path.join(os.path.expanduser("~"), "CNKI_检索报告.docx")
    print(f"\n保存路径（回车默认）: {default_path}")
    inp = input(">>> ").strip().strip('"').strip("'")
    file_path = inp if inp else default_path
    if not file_path.lower().endswith('.docx'):
        file_path += '.docx'
    print(f"   → {file_path}")

    # 启动 Chrome
    print("\n🔧 正在启动 Chrome 调试模式...")
    if not launch_chrome_debug():
        print("\n⚠ 自动启动失败，请手动启动 Chrome 调试模式：")
        print("   完全退出 Chrome 后运行：")
        if sys.platform == "darwin":
            print('   open -na "/Applications/Google Chrome.app" --args --remote-debugging-port=9222')
        else:
            print('   chrome --remote-debugging-port=9222')
        input("\n完成后按回车继续...")
        # 再检查一次
        if not launch_chrome_debug():
            print("❌ 无法连接 Chrome 调试端口")
            return

    print("\n🔑 请在打开的 Chrome 中：")
    print("   1) 访问 https://kns.cnki.net")
    print("   2) 登录知网账号")
    print("   3) 进行检索，进入结果页")
    input("\n完成后按回车继续...")

    print("⚡ 正在连接 Chrome...")
    try:
        driver = connect_chrome()
    except Exception as e:
        print(f"❌ 连接失败: {e}")
        return

    print("✅ 已连接，开始抓取...")
    papers = crawl_all_pages(driver)
    driver.quit()

    if not papers:
        print("❌ 未获取到任何论文")
        return

    print(f"\n✅ 共获取 {len(papers)} 篇论文，正在生成报告...")
    build_cnki_docx(papers, file_path)
    print(f"✅ 报告已保存到: {file_path}")


if __name__ == "__main__":
    if "--help" in sys.argv or "-h" in sys.argv:
        print_usage()
        sys.exit(0)
    cli_main()
